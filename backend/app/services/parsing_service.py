"""
Document parsing service.
Converts PDF/DOCX to canonical positioned blocks.
Phase 2 implementation.

FIX (2026-05-31):
  parse_docx now tracks the running section_heading (updated whenever a
  Heading-styled paragraph is encountered) and assigns it to every subsequent
  paragraph block.  This is the prerequisite for build_context_for_batch to
  filter blocks by semantic section — without it, every block has
  section_heading=None and hint matching degrades to checking only the
  first 120 chars of raw_text.
"""

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from typing import List, Dict, Optional
import uuid
from ..database import db_pool
from ..utils.bbox import normalize_bbox
import logging

logger = logging.getLogger(__name__)


class ParsingService:
    """Parses documents into canonical positioned blocks."""

    @staticmethod
    async def parse_pdf(file_path: str, contract_id: str) -> int:
        """
        Parse PDF into canonical blocks using PyMuPDF.

        Args:
            file_path: Path to PDF file
            contract_id: Contract UUID

        Returns:
            Number of blocks created
        """
        blocks_created = 0

        try:
            doc = fitz.open(file_path)
            blocks_to_insert = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                page_width = page.rect.width
                page_height = page.rect.height

                # Extract text blocks with positions
                blocks = page.get_text("dict")["blocks"]

                block_order = 0
                current_heading: Optional[str] = None

                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        # Extract text
                        lines = block.get("lines", [])
                        text_parts = []

                        for line in lines:
                            for span in line.get("spans", []):
                                text_parts.append(span.get("text", ""))

                        raw_text = " ".join(text_parts).strip()

                        if not raw_text:
                            continue

                        # Get bounding box
                        bbox = block.get("bbox", [0, 0, 0, 0])
                        x1, y1, x2, y2 = bbox

                        # Normalize coordinates
                        norm_x1, norm_y1, norm_x2, norm_y2 = normalize_bbox(
                            x1, y1, x2, y2, page_width, page_height
                        )

                        # Determine block type by font size heuristic
                        block_type = "paragraph"
                        max_font_size = 0.0
                        for line in lines:
                            for span in line.get("spans", []):
                                fs = span.get("size", 0)
                                if fs > max_font_size:
                                    max_font_size = fs

                        if max_font_size >= 13 or (len(raw_text) < 120 and raw_text.isupper()):
                            block_type = "heading"
                            current_heading = raw_text

                        blocks_to_insert.append({
                            "page_number":    page_num + 1,
                            "block_type":     block_type,
                            "raw_text":       raw_text,
                            "normalized_text": raw_text.lower().strip(),
                            "bbox_x1":        norm_x1,
                            "bbox_y1":        norm_y1,
                            "bbox_x2":        norm_x2,
                            "bbox_y2":        norm_y2,
                            "block_order":    block_order,
                            "section_heading": current_heading,
                            "table_context":  None
                        })

                        blocks_created += 1
                        block_order += 1

            if blocks_to_insert:
                await ParsingService._insert_blocks_bulk(contract_id, blocks_to_insert)

            num_pages = len(doc)  # capture before close() invalidates the object
            doc.close()
            logger.info(f"Parsed PDF: {blocks_created} blocks from {num_pages} pages")

            # Update contract page count
            await ParsingService._update_page_count(contract_id, num_pages)

            return blocks_created

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise

    @staticmethod
    async def parse_docx(file_path: str, contract_id: str) -> int:
        """
        Parse DOCX into canonical blocks.

        KEY FIX: Tracks the current section heading and assigns it to every
        paragraph block that follows it.  This is essential for
        build_context_for_batch to filter blocks by semantic section.

        Before this fix every block had section_heading=None, so hint matching
        fell back to checking only the first 120 chars of raw_text — causing
        Batch 1 (Metadata & Dates) to miss the preamble entirely on Indian MSAs
        where keywords like "effective date" don't appear in paragraph openers.

        Args:
            file_path: Path to DOCX file
            contract_id: Contract UUID

        Returns:
            Number of blocks created
        """
        blocks_created = 0

        try:
            doc = Document(file_path)
            block_order = 0
            current_heading: Optional[str] = None
            blocks_to_insert = []

            # Rough page tracking by word count
            word_total = 0
            _WORDS_PER_PAGE = 450

            for para in doc.paragraphs:
                raw_text = para.text.strip()

                if not raw_text:
                    continue

                # Estimate page number
                word_total += len(raw_text.split())
                page_num = max(1, (word_total // _WORDS_PER_PAGE) + 1)

                # Custom Heading Heuristics (Tier P3 Item 14):
                # Identify headings when native Word styles are omitted
                is_heading = False
                if para.style.name.startswith("Heading"):
                    is_heading = True
                elif len(raw_text) < 120 and raw_text.isupper():
                    is_heading = True
                elif len(raw_text) < 120:
                    # Heuristic: Check if the text is entirely bolded
                    non_empty_runs = [r for r in para.runs if r.text.strip()]
                    if non_empty_runs and all(r.bold for r in non_empty_runs):
                        is_heading = True

                if is_heading:
                    block_type = "heading"
                    current_heading = raw_text
                else:
                    block_type = "paragraph"

                # bbox_y encodes relative position so overlays have a distinct
                # y-coordinate even though DOCX has no real bounding boxes.
                blocks_to_insert.append({
                    "page_number":    page_num,
                    "block_type":     block_type,
                    "raw_text":       raw_text,
                    "normalized_text": raw_text.lower().strip(),
                    "bbox_x1":        0.0,
                    "bbox_y1":        round(block_order / 1000, 6),
                    "bbox_x2":        1.0,
                    "bbox_y2":        round((block_order + 1) / 1000, 6),
                    "block_order":    block_order,
                    "section_heading": current_heading,
                    "table_context":  None
                })

                blocks_created += 1
                block_order += 1

            if blocks_to_insert:
                await ParsingService._insert_blocks_bulk(contract_id, blocks_to_insert)

            logger.info(f"Parsed DOCX: {blocks_created} blocks")

            # Estimate page count
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            estimated_pages = max(1, word_count // _WORDS_PER_PAGE)
            await ParsingService._update_page_count(contract_id, estimated_pages)

            return blocks_created

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise

    @staticmethod
    async def _insert_blocks_bulk(contract_id: str, blocks: List[Dict]):
        """
        Bulk insert document blocks inside a single transaction with clean rollback (Tier P2 Item 9).
        Also wraps contract_id in HEXTORAW to prevent implicit conversion performance degradation.
        """
        if not blocks:
            return

        query = """
            INSERT INTO document_blocks (
                contract_id, page_number, block_type, raw_text, normalized_text,
                bbox_x1, bbox_y1, bbox_x2, bbox_y2, block_order,
                section_heading, table_context
            ) VALUES (
                HEXTORAW(:contract_id), :page_number, :block_type, :raw_text, :normalized_text,
                :bbox_x1, :bbox_y1, :bbox_x2, :bbox_y2, :block_order,
                :section_heading, :table_context
            )
        """

        # Prepare bind parameter dicts
        binds = []
        for b in blocks:
            binds.append({
                "contract_id":     contract_id,
                "page_number":     b["page_number"],
                "block_type":      b["block_type"],
                "raw_text":        b["raw_text"],
                "normalized_text": b["normalized_text"],
                "bbox_x1":         b["bbox_x1"],
                "bbox_y1":         b["bbox_y1"],
                "bbox_x2":         b["bbox_x2"],
                "bbox_y2":         b["bbox_y2"],
                "block_order":     b["block_order"],
                "section_heading": b["section_heading"],
                "table_context":   b["table_context"],
            })

        async with db_pool.get_connection() as conn:
            async with conn.cursor() as cursor:
                try:
                    await cursor.executemany(query, binds)
                    await conn.commit()
                    logger.info(f"Successfully bulk inserted {len(blocks)} blocks for contract {contract_id}")
                except Exception as err:
                    await conn.rollback()
                    logger.error(f"❌ Failed to bulk insert blocks for contract {contract_id}, rolling back: {err}")
                    raise

    @staticmethod
    async def _insert_block(
        contract_id: str,
        page_number: int,
        block_type: str,
        raw_text: str,
        normalized_text: str,
        bbox_x1: float,
        bbox_y1: float,
        bbox_x2: float,
        bbox_y2: float,
        block_order: int,
        section_heading: Optional[str] = None,
        table_context: Optional[str] = None,
    ):
        """Insert a canonical block into the database (backward compatible, wrapped in HEXTORAW)."""
        query = """
            INSERT INTO document_blocks (
                contract_id, page_number, block_type, raw_text, normalized_text,
                bbox_x1, bbox_y1, bbox_x2, bbox_y2, block_order,
                section_heading, table_context
            ) VALUES (
                HEXTORAW(:contract_id), :page_number, :block_type, :raw_text, :normalized_text,
                :bbox_x1, :bbox_y1, :bbox_x2, :bbox_y2, :block_order,
                :section_heading, :table_context
            )
        """

        async with db_pool.get_connection() as conn:
            async with conn.cursor() as cursor:
                await cursor.execute(query, {
                    "contract_id":    contract_id,
                    "page_number":    page_number,
                    "block_type":     block_type,
                    "raw_text":       raw_text,
                    "normalized_text": normalized_text,
                    "bbox_x1":        bbox_x1,
                    "bbox_y1":        bbox_y1,
                    "bbox_x2":        bbox_x2,
                    "bbox_y2":        bbox_y2,
                    "block_order":    block_order,
                    "section_heading": section_heading,
                    "table_context":  table_context,
                })
                await conn.commit()

    @staticmethod
    async def _update_page_count(contract_id: str, page_count: int):
        """Update contract page count (wrapped in HEXTORAW)."""
        query = """
            UPDATE contracts
            SET page_count = :page_count
            WHERE contract_id = HEXTORAW(:contract_id)
        """

        async with db_pool.get_connection() as conn:
            async with conn.cursor() as cursor:
                await cursor.execute(query, {
                    "contract_id": contract_id,
                    "page_count":  page_count,
                })
                await conn.commit()

    @staticmethod
    async def get_blocks_for_contract(contract_id: str) -> List[Dict]:
        """Retrieve all blocks for a contract (wrapped in HEXTORAW)."""
        query = """
            SELECT block_id, page_number, block_type, raw_text, normalized_text,
                   bbox_x1, bbox_y1, bbox_x2, bbox_y2, block_order,
                   section_heading, table_context
            FROM document_blocks
            WHERE contract_id = HEXTORAW(:contract_id)
            ORDER BY page_number, block_order
        """

        async with db_pool.get_connection() as conn:
            async with conn.cursor() as cursor:
                await cursor.execute(query, {"contract_id": contract_id})
                rows = await cursor.fetchall()

                blocks = []
                for row in rows:
                    raw_text = row[3]
                    if hasattr(raw_text, "read"):
                        raw_text = await raw_text.read()

                    normalized_text = row[4]
                    if hasattr(normalized_text, "read"):
                        normalized_text = await normalized_text.read()

                    blocks.append({
                        "block_id":       row[0],
                        "page_number":    row[1],
                        "block_type":     row[2],
                        "raw_text":       raw_text,
                        "normalized_text": normalized_text,
                        "bbox_x1":        row[5],
                        "bbox_y1":        row[6],
                        "bbox_x2":        row[7],
                        "bbox_y2":        row[8],
                        "block_order":    row[9],
                        "section_heading": row[10],
                        "table_context":  row[11],
                    })

                return blocks